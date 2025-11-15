"""
模板服务单元测试
测试模板上传、查询、删除、版本管理等功能
"""

import pytest
import tempfile
import shutil
from pathlib import Path
from io import BytesIO
from fastapi import UploadFile

from core.service.template_service import TemplateService
from core.models.template import Template


@pytest.fixture
def temp_dir():
    """创建临时目录"""
    temp_path = tempfile.mkdtemp()
    yield Path(temp_path)
    # 清理临时目录
    shutil.rmtree(temp_path, ignore_errors=True)


@pytest.fixture
def template_service(temp_dir):
    """创建模板服务实例"""
    from core.storage import TemplateStorage
    storage = TemplateStorage(base_path=str(temp_dir / "templates"))
    service = TemplateService(template_storage=storage)
    yield service


@pytest.fixture
def sample_html_content():
    """示例HTML模板内容"""
    return """
    <!DOCTYPE html>
    <html>
    <head>
        <title>{{ title }}</title>
    </head>
    <body>
        <h1>{{ heading }}</h1>
        <p>{{ content }}</p>
    </body>
    </html>
    """


@pytest.fixture
def sample_docx_file():
    """创建示例DOCX文件"""
    try:
        from docx import Document
        
        doc = Document()
        doc.add_heading('Test Template', 0)
        doc.add_paragraph('This is a test template.')
        
        # 保存到字节流
        file_stream = BytesIO()
        doc.save(file_stream)
        file_stream.seek(0)
        
        return file_stream
    except ImportError:
        pytest.skip("python-docx not installed")


async def create_upload_file(filename: str, content: bytes, content_type: str = "text/html"):
    """创建UploadFile对象"""
    file_stream = BytesIO(content)
    return UploadFile(
        filename=filename,
        file=file_stream,
        size=len(content),
        headers={"content-type": content_type}
    )


class TestTemplateServiceBasic:
    """模板服务基础功能测试"""
    
    @pytest.mark.asyncio
    async def test_create_html_template(self, template_service, sample_html_content):
        """测试创建HTML模板"""
        # 准备数据
        file = await create_upload_file(
            "test.html",
            sample_html_content.encode('utf-8'),
            "text/html"
        )
        metadata = {
            "name": "测试模板",
            "description": "这是一个测试模板",
            "tags": ["test", "html"],
            "version": "1.0.0",
        }
        
        # 创建模板
        template = await template_service.create_template(file, metadata)
        
        # 验证结果
        assert template.name == "测试模板"
        assert template.description == "这是一个测试模板"
        assert template.format == "html"
        assert template.version == "1.0.0"
        assert "test" in template.tags
        assert "html" in template.tags
        assert template.file_size > 0
        assert template.hash is not None
        assert template.template_id is not None
    
    @pytest.mark.asyncio
    async def test_create_docx_template(self, template_service, sample_docx_file):
        """测试创建DOCX模板"""
        # 准备数据
        content = sample_docx_file.read()
        sample_docx_file.seek(0)
        
        file = await create_upload_file(
            "test.docx",
            content,
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )
        metadata = {
            "name": "Word测试模板",
            "description": "这是一个Word测试模板",
            "tags": ["test", "docx"],
            "version": "1.0.0",
        }
        
        # 创建模板
        template = await template_service.create_template(file, metadata)
        
        # 验证结果
        assert template.name == "Word测试模板"
        assert template.format == "docx"
        assert template.version == "1.0.0"
        assert template.file_size > 0
    
    @pytest.mark.asyncio
    async def test_create_template_invalid_format(self, template_service):
        """测试创建不支持格式的模板"""
        # 准备数据
        file = await create_upload_file(
            "test.txt",
            b"test content",
            "text/plain"
        )
        metadata = {
            "name": "文本模板",
            "version": "1.0.0",
        }
        
        # 应该抛出异常
        with pytest.raises(ValueError, match="不支持的文件格式"):
            await template_service.create_template(file, metadata)
    
    @pytest.mark.asyncio
    async def test_get_template(self, template_service, sample_html_content):
        """测试获取模板"""
        # 先创建模板
        file = await create_upload_file("test.html", sample_html_content.encode('utf-8'))
        metadata = {"name": "测试模板", "version": "1.0.0"}
        created_template = await template_service.create_template(file, metadata)
        
        # 获取模板
        template = template_service.get_template(created_template.template_id)
        
        # 验证结果
        assert template.template_id == created_template.template_id
        assert template.name == created_template.name
        assert template.version == created_template.version
    
    def test_get_template_not_found(self, template_service):
        """测试获取不存在的模板"""
        with pytest.raises(FileNotFoundError):
            template_service.get_template("non_existent_id")
    
    @pytest.mark.asyncio
    async def test_list_templates_empty(self, template_service):
        """测试列表空模板"""
        result = template_service.list_templates({}, page=1, page_size=20)
        
        assert result["total"] == 0
        assert result["page"] == 1
        assert result["page_size"] == 20
        assert len(result["items"]) == 0
    
    @pytest.mark.asyncio
    async def test_list_templates(self, template_service, sample_html_content):
        """测试列表模板"""
        # 创建多个模板
        for i in range(5):
            file = await create_upload_file(f"test{i}.html", sample_html_content.encode('utf-8'))
            metadata = {
                "name": f"测试模板{i}",
                "version": "1.0.0",
                "tags": ["test"]
            }
            await template_service.create_template(file, metadata)
        
        # 列表查询
        result = template_service.list_templates({}, page=1, page_size=10)
        
        # 验证结果
        assert result["total"] == 5
        assert result["page"] == 1
        assert result["page_size"] == 10
        assert len(result["items"]) == 5
    
    @pytest.mark.asyncio
    async def test_list_templates_pagination(self, template_service, sample_html_content):
        """测试分页查询"""
        # 创建10个模板
        for i in range(10):
            file = await create_upload_file(f"test{i}.html", sample_html_content.encode('utf-8'))
            metadata = {"name": f"测试模板{i}", "version": "1.0.0"}
            await template_service.create_template(file, metadata)
        
        # 第一页
        result1 = template_service.list_templates({}, page=1, page_size=5)
        assert len(result1["items"]) == 5
        assert result1["total"] == 10
        
        # 第二页
        result2 = template_service.list_templates({}, page=2, page_size=5)
        assert len(result2["items"]) == 5
        assert result2["total"] == 10
        
        # 验证不重复
        ids1 = {item["template_id"] for item in result1["items"]}
        ids2 = {item["template_id"] for item in result2["items"]}
        assert len(ids1 & ids2) == 0  # 没有交集
    
    @pytest.mark.asyncio
    async def test_list_templates_filter_by_name(self, template_service, sample_html_content):
        """测试按名称筛选"""
        # 创建模板
        await template_service.create_template(
            await create_upload_file("test1.html", sample_html_content.encode('utf-8')),
            {"name": "销售报告", "version": "1.0.0"}
        )
        await template_service.create_template(
            await create_upload_file("test2.html", sample_html_content.encode('utf-8')),
            {"name": "财务报表", "version": "1.0.0"}
        )
        await template_service.create_template(
            await create_upload_file("test3.html", sample_html_content.encode('utf-8')),
            {"name": "销售统计", "version": "1.0.0"}
        )
        
        # 筛选包含"销售"的模板
        result = template_service.list_templates({"name": "销售"}, page=1, page_size=10)
        
        assert result["total"] == 2
        assert all("销售" in item["name"] for item in result["items"])
    
    @pytest.mark.asyncio
    async def test_list_templates_filter_by_tags(self, template_service, sample_html_content):
        """测试按标签筛选"""
        # 创建模板
        await template_service.create_template(
            await create_upload_file("test1.html", sample_html_content.encode('utf-8')),
            {"name": "模板1", "version": "1.0.0", "tags": ["report", "sales"]}
        )
        await template_service.create_template(
            await create_upload_file("test2.html", sample_html_content.encode('utf-8')),
            {"name": "模板2", "version": "1.0.0", "tags": ["report", "finance"]}
        )
        await template_service.create_template(
            await create_upload_file("test3.html", sample_html_content.encode('utf-8')),
            {"name": "模板3", "version": "1.0.0", "tags": ["invoice"]}
        )
        
        # 筛选包含"report"标签的模板
        result = template_service.list_templates({"tags": ["report"]}, page=1, page_size=10)
        
        assert result["total"] == 2
        assert all("report" in item["tags"] for item in result["items"])
    
    @pytest.mark.asyncio
    async def test_list_templates_filter_by_format(self, template_service, sample_html_content, sample_docx_file):
        """测试按格式筛选"""
        # 创建HTML模板
        await template_service.create_template(
            await create_upload_file("test1.html", sample_html_content.encode('utf-8')),
            {"name": "HTML模板", "version": "1.0.0"}
        )
        
        # 创建DOCX模板
        content = sample_docx_file.read()
        await template_service.create_template(
            await create_upload_file(
                "test.docx",
                content,
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            ),
            {"name": "Word模板", "version": "1.0.0"}
        )
        
        # 筛选HTML格式
        result = template_service.list_templates({"format": "html"}, page=1, page_size=10)
        assert result["total"] == 1
        assert result["items"][0]["format"] == "html"
        
        # 筛选DOCX格式
        result = template_service.list_templates({"format": "docx"}, page=1, page_size=10)
        assert result["total"] == 1
        assert result["items"][0]["format"] == "docx"


class TestTemplateServiceUpdate:
    """模板服务更新功能测试"""
    
    @pytest.mark.asyncio
    async def test_update_template_name(self, template_service, sample_html_content):
        """测试更新模板名称"""
        # 创建模板
        file = await create_upload_file("test.html", sample_html_content.encode('utf-8'))
        template = await template_service.create_template(
            file,
            {"name": "旧名称", "version": "1.0.0"}
        )
        
        # 更新名称
        updated = template_service.update_template(
            template.template_id,
            {"name": "新名称"}
        )
        
        assert updated.name == "新名称"
        assert updated.template_id == template.template_id
    
    @pytest.mark.asyncio
    async def test_update_template_description(self, template_service, sample_html_content):
        """测试更新模板描述"""
        # 创建模板
        file = await create_upload_file("test.html", sample_html_content.encode('utf-8'))
        template = await template_service.create_template(
            file,
            {"name": "测试", "version": "1.0.0", "description": "旧描述"}
        )
        
        # 更新描述
        updated = template_service.update_template(
            template.template_id,
            {"description": "新描述"}
        )
        
        assert updated.description == "新描述"
    
    @pytest.mark.asyncio
    async def test_update_template_tags(self, template_service, sample_html_content):
        """测试更新模板标签"""
        # 创建模板
        file = await create_upload_file("test.html", sample_html_content.encode('utf-8'))
        template = await template_service.create_template(
            file,
            {"name": "测试", "version": "1.0.0", "tags": ["old"]}
        )
        
        # 更新标签
        updated = template_service.update_template(
            template.template_id,
            {"tags": ["new", "updated"]}
        )
        
        assert "new" in updated.tags
        assert "updated" in updated.tags
        assert "old" not in updated.tags
    
    def test_update_template_not_found(self, template_service):
        """测试更新不存在的模板"""
        with pytest.raises(FileNotFoundError):
            template_service.update_template("non_existent_id", {"name": "新名称"})


class TestTemplateServiceDelete:
    """模板服务删除功能测试"""
    
    @pytest.mark.asyncio
    async def test_delete_template(self, template_service, sample_html_content):
        """测试删除模板"""
        # 创建模板
        file = await create_upload_file("test.html", sample_html_content.encode('utf-8'))
        template = await template_service.create_template(
            file,
            {"name": "测试模板", "version": "1.0.0"}
        )
        
        # 删除模板
        success = template_service.delete_template(template.template_id)
        assert success is True
        
        # 验证已删除
        with pytest.raises(FileNotFoundError):
            template_service.get_template(template.template_id)
    
    def test_delete_template_not_found(self, template_service):
        """测试删除不存在的模板"""
        with pytest.raises(FileNotFoundError):
            template_service.delete_template("non_existent_id")


class TestTemplateServiceVersions:
    """模板服务版本管理测试"""
    
    @pytest.mark.asyncio
    async def test_create_version(self, template_service, sample_html_content):
        """测试创建新版本"""
        # 创建初始模板
        file = await create_upload_file("test.html", sample_html_content.encode('utf-8'))
        template = await template_service.create_template(
            file,
            {"name": "测试模板", "version": "1.0.0"}
        )
        
        # 创建新版本
        new_content = sample_html_content.replace("{{ title }}", "{{ new_title }}")
        new_file = await create_upload_file("test.html", new_content.encode('utf-8'))
        
        version = await template_service.create_version(
            template_id=template.template_id,
            file=new_file,
            version="1.1.0",
            changelog="更新标题字段"
        )
        
        # 验证结果
        assert version.template_id == template.template_id
        assert version.version == "1.1.0"
        assert version.changelog == "更新标题字段"
        assert version.hash != template.hash  # 内容不同，哈希应该不同
    
    @pytest.mark.asyncio
    async def test_create_version_not_found(self, template_service, sample_html_content):
        """测试为不存在的模板创建版本"""
        file = await create_upload_file("test.html", sample_html_content.encode('utf-8'))
        
        with pytest.raises(FileNotFoundError):
            await template_service.create_version(
                template_id="non_existent_id",
                file=file,
                version="2.0.0",
                changelog="测试"
            )
    
    @pytest.mark.asyncio
    async def test_list_versions(self, template_service, sample_html_content):
        """测试列表版本"""
        # 创建初始模板
        file = await create_upload_file("test.html", sample_html_content.encode('utf-8'))
        template = await template_service.create_template(
            file,
            {"name": "测试模板", "version": "1.0.0"}
        )
        
        # 创建多个版本
        for i in range(1, 4):
            new_file = await create_upload_file("test.html", sample_html_content.encode('utf-8'))
            await template_service.create_version(
                template_id=template.template_id,
                file=new_file,
                version=f"1.{i}.0",
                changelog=f"版本 1.{i}.0"
            )
        
        # 列表版本
        result = template_service.list_versions(template.template_id, page=1, page_size=10)
        
        # 验证结果（包括初始版本，总共4个）
        assert result["total"] == 4
        assert len(result["items"]) == 4
        
        # 验证版本排序（最新的在前）
        versions = [item["version"] for item in result["items"]]
        assert versions[0] == "1.3.0"
        assert versions[-1] == "1.0.0"
    
    @pytest.mark.asyncio
    async def test_get_specific_version(self, template_service, sample_html_content):
        """测试获取特定版本"""
        # 创建初始模板
        file = await create_upload_file("test.html", sample_html_content.encode('utf-8'))
        template = await template_service.create_template(
            file,
            {"name": "测试模板", "version": "1.0.0"}
        )
        
        # 创建新版本
        new_file = await create_upload_file("test.html", b"new content")
        await template_service.create_version(
            template_id=template.template_id,
            file=new_file,
            version="2.0.0",
            changelog="大版本更新"
        )
        
        # 获取1.0.0版本
        v1 = template_service.get_template(template.template_id, version="1.0.0")
        assert v1.version == "1.0.0"
        
        # 获取2.0.0版本
        v2 = template_service.get_template(template.template_id, version="2.0.0")
        assert v2.version == "2.0.0"
        
        # 获取最新版本（不指定版本号）
        latest = template_service.get_template(template.template_id)
        assert latest.version == "2.0.0"
    
    @pytest.mark.asyncio
    async def test_delete_specific_version(self, template_service, sample_html_content):
        """测试删除特定版本"""
        # 创建初始模板
        file = await create_upload_file("test.html", sample_html_content.encode('utf-8'))
        template = await template_service.create_template(
            file,
            {"name": "测试模板", "version": "1.0.0"}
        )
        
        # 创建新版本
        new_file = await create_upload_file("test.html", b"new content")
        await template_service.create_version(
            template_id=template.template_id,
            file=new_file,
            version="2.0.0",
            changelog="新版本"
        )
        
        # 删除1.0.0版本
        success = template_service.delete_template(template.template_id, version="1.0.0")
        assert success is True
        
        # 验证1.0.0版本已删除
        with pytest.raises(FileNotFoundError):
            template_service.get_template(template.template_id, version="1.0.0")
        
        # 验证2.0.0版本仍然存在
        v2 = template_service.get_template(template.template_id, version="2.0.0")
        assert v2.version == "2.0.0"


class TestTemplateServiceDownload:
    """模板服务下载功能测试"""
    
    @pytest.mark.asyncio
    async def test_download_template(self, template_service, sample_html_content):
        """测试下载模板"""
        # 创建模板
        file = await create_upload_file("test.html", sample_html_content.encode('utf-8'))
        template = await template_service.create_template(
            file,
            {"name": "测试模板", "version": "1.0.0"}
        )
        
        # 下载模板
        content = template_service.download_template(template.template_id)
        
        # 验证内容
        assert content is not None
        assert len(content) > 0
        assert content == sample_html_content.encode('utf-8')
    
    @pytest.mark.asyncio
    async def test_download_specific_version(self, template_service, sample_html_content):
        """测试下载特定版本"""
        # 创建初始模板
        file = await create_upload_file("test.html", sample_html_content.encode('utf-8'))
        template = await template_service.create_template(
            file,
            {"name": "测试模板", "version": "1.0.0"}
        )
        
        # 创建新版本
        new_content = b"new version content"
        new_file = await create_upload_file("test.html", new_content)
        await template_service.create_version(
            template_id=template.template_id,
            file=new_file,
            version="2.0.0",
            changelog="新版本"
        )
        
        # 下载1.0.0版本
        content_v1 = template_service.download_template(template.template_id, version="1.0.0")
        assert content_v1 == sample_html_content.encode('utf-8')
        
        # 下载2.0.0版本
        content_v2 = template_service.download_template(template.template_id, version="2.0.0")
        assert content_v2 == new_content
        
        # 下载最新版本
        content_latest = template_service.download_template(template.template_id)
        assert content_latest == new_content
    
    def test_download_template_not_found(self, template_service):
        """测试下载不存在的模板"""
        with pytest.raises(FileNotFoundError):
            template_service.download_template("non_existent_id")

