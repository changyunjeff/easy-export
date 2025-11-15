#!/usr/bin/env python
# -*- coding: utf-8 -*-
"""
文档校验 MVP
用于验证生成的文档是否满足要求

功能：
1. 必填字段检查 - 检查文档中是否包含所有必填字段
2. 数据对齐检查 - 检查表格行数与数据是否匹配
3. 链接有效性检查 - 检查文档中的链接是否可访问（可选）
4. 样式一致性检查 - 检查字体、颜色等样式是否一致（可选）

使用方法：
    python mvp/validate_document.py output.html mvp/validation_rules.json
    python mvp/validate_document.py output.docx mvp/validation_rules.json --verbose
"""

import sys
import json
import argparse
import re
from pathlib import Path
from typing import Dict, List, Any, Optional
from dataclasses import dataclass, asdict
import logging

# 配置日志
logging.basicConfig(level=logging.INFO, format='%(levelname)s: %(message)s')
logger = logging.getLogger(__name__)


@dataclass
class ValidationError:
    """校验错误"""
    type: str
    message: str
    field: Optional[str] = None
    detail: Optional[Dict[str, Any]] = None


@dataclass
class ValidationWarning:
    """校验警告"""
    type: str
    message: str
    detail: Optional[Dict[str, Any]] = None


@dataclass
class ValidationResult:
    """校验结果"""
    passed: bool
    errors: List[Dict[str, Any]]
    warnings: List[Dict[str, Any]]
    summary: Dict[str, Any]


class DocumentValidator:
    """文档校验器"""
    
    def __init__(self):
        """初始化校验器"""
        self.errors: List[ValidationError] = []
        self.warnings: List[ValidationWarning] = []
    
    def validate(
        self,
        file_path: str,
        rules: Dict[str, Any]
    ) -> ValidationResult:
        """
        校验文档
        
        Args:
            file_path: 文档路径
            rules: 校验规则
                - required_fields: 必填字段列表
                - check_links: 是否检查链接（默认False）
                - check_style: 是否检查样式（默认False）
                - expected_table_rows: 预期表格行数（可选）
                - link_timeout_sec: 链接超时时间（秒，默认3）
        
        Returns:
            ValidationResult: 校验结果
        """
        self.errors = []
        self.warnings = []
        
        file_path_obj = Path(file_path)
        
        # 检查文件是否存在
        if not file_path_obj.exists():
            self.errors.append(ValidationError(
                type="file_not_found",
                message=f"文件不存在: {file_path}"
            ))
            return self._build_result()
        
        # 读取文件内容
        try:
            if file_path_obj.suffix.lower() in ['.html', '.htm']:
                content = self._read_html(file_path)
            elif file_path_obj.suffix.lower() == '.docx':
                content = self._read_docx(file_path)
            else:
                self.errors.append(ValidationError(
                    type="unsupported_format",
                    message=f"不支持的文件格式: {file_path_obj.suffix}"
                ))
                return self._build_result()
        except Exception as e:
            self.errors.append(ValidationError(
                type="read_error",
                message=f"读取文件失败: {str(e)}"
            ))
            return self._build_result()
        
        # 1. 必填字段检查
        if rules.get("required_fields"):
            self._check_required_fields(content, rules["required_fields"])
        
        # 2. 表格行数检查
        if "expected_table_rows" in rules:
            self._check_table_rows(content, rules["expected_table_rows"])
        
        # 3. 链接有效性检查（可选）
        if rules.get("check_links", False):
            timeout = rules.get("link_timeout_sec", 3)
            self._check_links(content, timeout)
        
        # 4. 样式一致性检查（可选）
        if rules.get("check_style", False):
            self._check_style_consistency(content)
        
        return self._build_result()
    
    def _read_html(self, file_path: str) -> str:
        """读取HTML文件"""
        with open(file_path, 'r', encoding='utf-8') as f:
            return f.read()
    
    def _read_docx(self, file_path: str) -> str:
        """读取Word文档内容"""
        try:
            from docx import Document
            doc = Document(file_path)
            # 提取所有段落和表格文本
            content = []
            for para in doc.paragraphs:
                content.append(para.text)
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        content.append(cell.text)
            return '\n'.join(content)
        except ImportError:
            raise ImportError("需要安装python-docx: pip install python-docx")
    
    def _check_required_fields(
        self,
        content: str,
        required_fields: List[str]
    ) -> None:
        """检查必填字段"""
        logger.info(f"检查必填字段: {required_fields}")
        
        missing_fields = []
        for field in required_fields:
            # 简单的字符串匹配检查
            if field not in content:
                missing_fields.append(field)
        
        if missing_fields:
            for field in missing_fields:
                self.errors.append(ValidationError(
                    type="missing_required_field",
                    message=f"缺失必填字段: {field}",
                    field=field
                ))
            logger.error(f"发现 {len(missing_fields)} 个缺失的必填字段")
        else:
            logger.info("✓ 所有必填字段都存在")
    
    def _check_table_rows(
        self,
        content: str,
        expected_rows: int
    ) -> None:
        """检查表格行数"""
        logger.info(f"检查表格行数，预期: {expected_rows}")
        
        # HTML表格行数检查
        if '<table' in content.lower():
            # 简单统计<tr>标签数量
            tr_count = content.lower().count('<tr')
            # 通常第一行是表头，所以数据行 = tr_count - 1
            data_rows = max(0, tr_count - 1)
            
            if data_rows != expected_rows:
                self.warnings.append(ValidationWarning(
                    type="table_row_mismatch",
                    message=f"表格行数不匹配: 预期 {expected_rows} 行，实际 {data_rows} 行",
                    detail={
                        "expected": expected_rows,
                        "actual": data_rows
                    }
                ))
                logger.warning(f"⚠ 表格行数不匹配")
            else:
                logger.info(f"✓ 表格行数匹配 ({data_rows} 行)")
        else:
            logger.info("未找到表格，跳过行数检查")
    
    def _check_links(
        self,
        content: str,
        timeout: int = 3
    ) -> None:
        """检查链接有效性"""
        logger.info(f"检查链接有效性 (超时: {timeout}秒)")
        
        # 提取所有HTTP/HTTPS链接
        url_pattern = r'https?://[^\s<>"]+|www\.[^\s<>"]+'
        urls = re.findall(url_pattern, content)
        
        if not urls:
            logger.info("未找到链接，跳过检查")
            return
        
        logger.info(f"找到 {len(urls)} 个链接")
        
        # 检查每个链接（需要requests库）
        try:
            import requests
            
            broken_links = []
            slow_links = []
            
            for url in urls:
                try:
                    response = requests.head(url, timeout=timeout, allow_redirects=True)
                    if response.status_code >= 400:
                        broken_links.append({
                            "url": url,
                            "status_code": response.status_code
                        })
                    elif response.elapsed.total_seconds() > timeout * 0.8:
                        slow_links.append({
                            "url": url,
                            "elapsed": response.elapsed.total_seconds()
                        })
                except Exception as e:
                    broken_links.append({
                        "url": url,
                        "error": str(e)
                    })
            
            # 记录错误和警告
            for link in broken_links:
                self.errors.append(ValidationError(
                    type="broken_link",
                    message=f"链接无法访问: {link['url']}",
                    detail=link
                ))
            
            for link in slow_links:
                self.warnings.append(ValidationWarning(
                    type="slow_link",
                    message=f"链接响应较慢: {link['url']} ({link['elapsed']:.2f}秒)",
                    detail=link
                ))
            
            if broken_links:
                logger.error(f"发现 {len(broken_links)} 个无效链接")
            if slow_links:
                logger.warning(f"发现 {len(slow_links)} 个响应慢的链接")
            if not broken_links and not slow_links:
                logger.info("✓ 所有链接都正常")
        
        except ImportError:
            logger.warning("未安装requests库，跳过链接检查")
    
    def _check_style_consistency(self, content: str) -> None:
        """检查样式一致性（简化实现）"""
        logger.info("检查样式一致性")
        
        # HTML样式检查
        if '<style' in content.lower() or 'style=' in content.lower():
            # 提取所有内联样式
            inline_styles = re.findall(r'style="([^"]+)"', content, re.IGNORECASE)
            
            # 检查字体是否一致
            font_families = set()
            for style in inline_styles:
                font_match = re.search(r'font-family:\s*([^;]+)', style, re.IGNORECASE)
                if font_match:
                    font_families.add(font_match.group(1).strip())
            
            if len(font_families) > 3:
                self.warnings.append(ValidationWarning(
                    type="inconsistent_fonts",
                    message=f"文档中使用了过多的字体 ({len(font_families)} 种)",
                    detail={"fonts": list(font_families)}
                ))
                logger.warning(f"⚠ 使用了 {len(font_families)} 种字体，建议统一")
            else:
                logger.info(f"✓ 字体使用适当 ({len(font_families)} 种)")
        else:
            logger.info("未找到样式信息，跳过检查")
    
    def _build_result(self) -> ValidationResult:
        """构建校验结果"""
        passed = len(self.errors) == 0
        
        return ValidationResult(
            passed=passed,
            errors=[asdict(e) for e in self.errors],
            warnings=[asdict(w) for w in self.warnings],
            summary={
                "total_checks": len(self.errors) + len(self.warnings),
                "passed_checks": 0 if self.errors else len(self.warnings),
                "failed_checks": len(self.errors),
                "warning_checks": len(self.warnings)
            }
        )


def main():
    """主函数"""
    parser = argparse.ArgumentParser(description='文档校验工具')
    parser.add_argument('file_path', help='要校验的文档路径')
    parser.add_argument('rules_file', help='校验规则JSON文件路径')
    parser.add_argument('--verbose', '-v', action='store_true', help='详细输出')
    parser.add_argument('--output', '-o', help='输出报告JSON文件路径')
    
    args = parser.parse_args()
    
    if args.verbose:
        logging.getLogger().setLevel(logging.DEBUG)
    
    # 读取校验规则
    try:
        with open(args.rules_file, 'r', encoding='utf-8') as f:
            rules = json.load(f)
    except Exception as e:
        logger.error(f"读取规则文件失败: {e}")
        sys.exit(1)
    
    # 执行校验
    validator = DocumentValidator()
    result = validator.validate(args.file_path, rules)
    
    # 输出结果
    print("\n" + "="*60)
    print("文档校验结果")
    print("="*60)
    print(f"文件: {args.file_path}")
    print(f"状态: {'✓ 通过' if result.passed else '✗ 失败'}")
    print(f"错误: {len(result.errors)} 个")
    print(f"警告: {len(result.warnings)} 个")
    
    if result.errors:
        print("\n错误详情:")
        for error in result.errors:
            print(f"  - [{error['type']}] {error['message']}")
            if error.get('detail'):
                print(f"    详情: {error['detail']}")
    
    if result.warnings:
        print("\n警告详情:")
        for warning in result.warnings:
            print(f"  - [{warning['type']}] {warning['message']}")
            if warning.get('detail'):
                print(f"    详情: {warning['detail']}")
    
    print("\n汇总:")
    print(f"  总检查项: {result.summary['total_checks']}")
    print(f"  通过检查: {result.summary['passed_checks']}")
    print(f"  失败检查: {result.summary['failed_checks']}")
    print(f"  警告检查: {result.summary['warning_checks']}")
    print("="*60)
    
    # 保存报告
    if args.output:
        output_data = asdict(result)
        with open(args.output, 'w', encoding='utf-8') as f:
            json.dump(output_data, f, ensure_ascii=False, indent=2)
        logger.info(f"校验报告已保存到: {args.output}")
    
    # 返回退出码
    sys.exit(0 if result.passed else 1)


if __name__ == "__main__":
    main()

