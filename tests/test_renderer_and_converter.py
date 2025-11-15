"""
测试渲染器和格式转换器
"""

import os
import sys

# 确保MSYS2的bin目录在PATH最前面
msys2_bin = r'D:\chang\app\msys2\ucrt64\bin'
if msys2_bin not in os.environ['PATH'].split(os.pathsep)[0]:
    os.environ['PATH'] = msys2_bin + os.pathsep + os.environ['PATH']

if sys.version_info >= (3, 8):
    os.add_dll_directory(msys2_bin)

import io
import zipfile
from pathlib import Path

import pytest

from core.engine.renderer import (
    DocxRenderer,
    HTMLRenderer,
    PDFRenderer,
    RendererFactory,
)
from core.engine.converter import Converter
from core.engine.template import Template


@pytest.fixture
def html_template() -> Template:
    """HTML模板示例"""
    content = """
    <!DOCTYPE html>
    <html>
    <head>
        <title>{{ title }}</title>
    </head>
    <body>
        <h1>{{ heading }}</h1>
        <p>{{ description }}</p>
        <ul>
        {% for item in items %}
            <li>{{ item }}</li>
        {% endfor %}
        </ul>
    </body>
    </html>
    """.strip().encode("utf-8")
    
    return Template(
        template_id="tpl_html_test",
        version="v1.0.0",
        format="html",
        content=content,
        placeholders=None,
    )


@pytest.fixture
def docx_template() -> Template:
    """Word模板示例（使用docxtpl语法）"""
    # 创建一个简单的docx文件，包含Jinja2占位符
    try:
        from docx import Document
        from docxtpl import DocxTemplate
        import tempfile
        import os
        
        # 先创建一个基础的docx文件
        with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as temp:
            temp_path = temp.name
        
        # 使用python-docx创建基础文档
        doc = Document()
        doc.add_paragraph('{{ title }}')
        doc.add_paragraph('{{ description }}')
        doc.save(temp_path)
        
        # 读取内容
        with open(temp_path, 'rb') as f:
            content = f.read()
        
        # 清理临时文件
        try:
            os.unlink(temp_path)
        except:
            pass
        
        return Template(
            template_id="tpl_docx_test",
            version="v1.0.0",
            format="docx",
            content=content,
            placeholders=None,
        )
    except ImportError:
        pytest.skip("需要安装docxtpl库")


@pytest.fixture
def sample_data() -> dict:
    """示例数据"""
    return {
        "title": "测试标题",
        "heading": "主标题",
        "description": "这是一段描述文字",
        "items": ["项目1", "项目2", "项目3"],
    }


# =============================================================================
# HTMLRenderer 测试
# =============================================================================

def test_html_renderer_basic(html_template: Template, sample_data: dict) -> None:
    """测试HTML渲染器基础功能"""
    renderer = HTMLRenderer()
    result = renderer.render(html_template, sample_data)
    
    assert isinstance(result, bytes)
    content = result.decode("utf-8")
    
    assert "测试标题" in content
    assert "主标题" in content
    assert "这是一段描述文字" in content
    assert "项目1" in content
    assert "项目2" in content
    assert "项目3" in content


def test_html_renderer_supports_format() -> None:
    """测试HTML渲染器格式支持"""
    renderer = HTMLRenderer()
    
    assert renderer.supports_format("html") is True
    assert renderer.supports_format("htm") is True
    assert renderer.supports_format("HTML") is True
    assert renderer.supports_format("pdf") is False
    assert renderer.supports_format("docx") is False


def test_html_renderer_empty_data(html_template: Template) -> None:
    """测试HTML渲染器处理空数据"""
    renderer = HTMLRenderer()
    
    # 空数据应该仍能渲染，只是占位符为空或使用默认值
    result = renderer.render(html_template, {})
    assert isinstance(result, bytes)


# =============================================================================
# DocxRenderer 测试
# =============================================================================

def test_docx_renderer_basic(docx_template: Template, sample_data: dict) -> None:
    """测试Word渲染器基础功能"""
    try:
        from docx import Document
    except ImportError:
        pytest.skip("需要安装python-docx库")
    
    renderer = DocxRenderer()
    result = renderer.render(docx_template, sample_data)
    
    assert isinstance(result, bytes)
    assert len(result) > 0
    
    # 验证生成的docx文件可以被读取
    doc = Document(io.BytesIO(result))
    text = "\n".join([paragraph.text for paragraph in doc.paragraphs])
    
    assert "测试标题" in text
    assert "这是一段描述文字" in text


def test_docx_renderer_supports_format() -> None:
    """测试Word渲染器格式支持"""
    renderer = DocxRenderer()
    
    assert renderer.supports_format("docx") is True
    assert renderer.supports_format("DOCX") is True
    assert renderer.supports_format("html") is False
    assert renderer.supports_format("pdf") is False


def test_docx_renderer_missing_dependency() -> None:
    """测试Word渲染器缺少依赖的情况"""
    # 这个测试假设docxtpl已安装，如果未安装会抛出ImportError
    renderer = DocxRenderer()
    
    # 创建一个简单的模板
    template = Template(
        template_id="test",
        version="v1",
        format="docx",
        content=b"fake content",
        placeholders=None,
    )
    
    # 如果docxtpl未安装，应该抛出ImportError
    # 如果已安装，会因为内容无效而抛出其他错误
    with pytest.raises((ImportError, ValueError)):
        renderer.render(template, {})


# =============================================================================
# Converter 测试
# =============================================================================

def test_converter_html_to_pdf_basic() -> None:
    """测试HTML转PDF基础功能"""
    try:
        from weasyprint import HTML
    except (ImportError, OSError) as e:
        pytest.skip(f"WeasyPrint不可用: {e}")
    
    converter = Converter()
    
    html = """
    <!DOCTYPE html>
    <html>
    <head><title>Test</title></head>
    <body>
        <h1>测试标题</h1>
        <p>测试内容</p>
    </body>
    </html>
    """
    
    pdf_bytes = converter.html_to_pdf(html)
    
    assert isinstance(pdf_bytes, bytes)
    assert len(pdf_bytes) > 0
    # PDF文件应该以%PDF开头
    assert pdf_bytes.startswith(b"%PDF")


def test_converter_html_to_pdf_with_css() -> None:
    """测试HTML转PDF带CSS样式"""
    try:
        from weasyprint import HTML
    except (ImportError, OSError) as e:
        pytest.skip(f"WeasyPrint不可用: {e}")
    
    converter = Converter()
    
    html = """
    <!DOCTYPE html>
    <html>
    <body>
        <h1 class="title">标题</h1>
    </body>
    </html>
    """
    
    css = """
    .title {
        color: red;
        font-size: 24px;
    }
    """
    
    pdf_bytes = converter.html_to_pdf(html, css)
    
    assert isinstance(pdf_bytes, bytes)
    assert len(pdf_bytes) > 0
    assert pdf_bytes.startswith(b"%PDF")


def test_converter_docx_to_pdf_basic(docx_template: Template) -> None:
    """测试Word转PDF基础功能"""
    try:
        from docx2pdf import convert
    except ImportError:
        pytest.skip("需要安装docx2pdf库")
    
    converter = Converter()
    
    # 使用docx_template的内容
    pdf_bytes = converter.docx_to_pdf(docx_template.content)
    
    assert isinstance(pdf_bytes, bytes)
    assert len(pdf_bytes) > 0
    # PDF文件应该以%PDF开头
    assert pdf_bytes.startswith(b"%PDF")


def test_converter_pdf_to_html_not_implemented() -> None:
    """测试PDF转HTML功能未实现"""
    converter = Converter()
    
    with pytest.raises(NotImplementedError):
        converter.pdf_to_html(b"fake pdf content")


# =============================================================================
# PDFRenderer 测试
# =============================================================================

def test_pdf_renderer_from_html_template(html_template: Template, sample_data: dict) -> None:
    """测试从HTML模板渲染PDF"""
    try:
        from weasyprint import HTML
    except (ImportError, OSError) as e:
        pytest.skip(f"WeasyPrint不可用: {e}")
    
    renderer = PDFRenderer()
    result = renderer.render(html_template, sample_data)
    
    assert isinstance(result, bytes)
    assert len(result) > 0
    assert result.startswith(b"%PDF")


def test_pdf_renderer_from_docx_template(docx_template: Template, sample_data: dict) -> None:
    """测试从Word模板渲染PDF"""
    try:
        from docx2pdf import convert
    except ImportError:
        pytest.skip("需要安装docx2pdf库")
    
    renderer = PDFRenderer()
    result = renderer.render(docx_template, sample_data)
    
    assert isinstance(result, bytes)
    assert len(result) > 0
    assert result.startswith(b"%PDF")


def test_pdf_renderer_supports_format() -> None:
    """测试PDF渲染器格式支持"""
    renderer = PDFRenderer()
    
    assert renderer.supports_format("pdf") is True
    assert renderer.supports_format("PDF") is True
    assert renderer.supports_format("html") is False
    assert renderer.supports_format("docx") is False


def test_pdf_renderer_supports_template() -> None:
    """测试PDF渲染器模板格式支持"""
    renderer = PDFRenderer()
    
    # PDF渲染器应该支持HTML和docx模板
    html_tpl = Template("id", "v1", "html", b"", None)
    docx_tpl = Template("id", "v1", "docx", b"", None)
    pdf_tpl = Template("id", "v1", "pdf", b"", None)
    
    assert renderer.supports_template(html_tpl) is True
    assert renderer.supports_template(docx_tpl) is True
    assert renderer.supports_template(pdf_tpl) is False


# =============================================================================
# RendererFactory 测试
# =============================================================================

def test_renderer_factory_get_html_renderer() -> None:
    """测试获取HTML渲染器"""
    renderer = RendererFactory.get_renderer("html")
    
    assert isinstance(renderer, HTMLRenderer)
    assert renderer.supports_format("html") is True


def test_renderer_factory_get_docx_renderer() -> None:
    """测试获取Word渲染器"""
    renderer = RendererFactory.get_renderer("docx")
    
    assert isinstance(renderer, DocxRenderer)
    assert renderer.supports_format("docx") is True


def test_renderer_factory_get_pdf_renderer() -> None:
    """测试获取PDF渲染器"""
    renderer = RendererFactory.get_renderer("pdf")
    
    assert isinstance(renderer, PDFRenderer)
    assert renderer.supports_format("pdf") is True


def test_renderer_factory_unsupported_format() -> None:
    """测试不支持的格式"""
    with pytest.raises(ValueError, match="Unsupported format"):
        RendererFactory.get_renderer("unknown")


def test_renderer_factory_case_insensitive() -> None:
    """测试格式名大小写不敏感"""
    renderer1 = RendererFactory.get_renderer("HTML")
    renderer2 = RendererFactory.get_renderer("html")
    renderer3 = RendererFactory.get_renderer("Html")
    
    # 应该返回相同的渲染器实例
    assert renderer1 is renderer2
    assert renderer2 is renderer3

