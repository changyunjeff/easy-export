"""
校验服务模块
负责文档格式校验、数据完整性检查、链接有效性验证
"""

import logging
import re
from pathlib import Path
from typing import List, Dict, Any, Optional

logger = logging.getLogger(__name__)


class ValidationError:
    """校验错误类"""
    
    def __init__(self, type: str, message: str, field: Optional[str] = None, detail: Dict[str, Any] = None):
        """
        初始化校验错误
        
        Args:
            type: 错误类型
            message: 错误消息
            field: 相关字段（可选）
            detail: 错误详情
        """
        self.type = type
        self.message = message
        self.field = field
        self.detail = detail or {}
    
    def to_dict(self) -> Dict[str, Any]:
        """转换为字典"""
        result = {
            "type": self.type,
            "message": self.message
        }
        if self.field:
            result["field"] = self.field
        if self.detail:
            result["detail"] = self.detail
        return result


class ValidationWarning:
    """校验警告类"""
    
    def __init__(self, type: str, message: str, detail: Dict[str, Any] = None):
        """
        初始化校验警告
        
        Args:
            type: 警告类型
            message: 警告消息
            detail: 警告详情
        """
        self.type = type
        self.message = message
        self.detail = detail or {}
    
    def to_dict(self) -> Dict[str, Any]:
        """转换为字典"""
        result = {
            "type": self.type,
            "message": self.message
        }
        if self.detail:
            result["detail"] = self.detail
        return result


class ValidationResult:
    """校验结果类"""
    
    def __init__(
        self,
        passed: bool,
        errors: List[ValidationError] = None,
        warnings: List[ValidationWarning] = None,
    ):
        """
        初始化校验结果
        
        Args:
            passed: 是否通过
            errors: 错误列表
            warnings: 警告列表
        """
        self.passed = passed
        self.errors = errors or []
        self.warnings = warnings or []
    
    def to_dict(self) -> Dict[str, Any]:
        """转换为字典"""
        return {
            "passed": self.passed,
            "errors": [e.to_dict() for e in self.errors],
            "warnings": [w.to_dict() for w in self.warnings],
            "summary": {
                "total_checks": len(self.errors) + len(self.warnings),
                "passed_checks": 0 if self.errors else len(self.warnings),
                "failed_checks": len(self.errors),
                "warning_checks": len(self.warnings)
            }
        }


class ValidateService:
    """
    校验服务类
    负责文档格式校验、数据完整性检查等
    """
    
    def __init__(self):
        """初始化校验服务"""
        self.errors: List[ValidationError] = []
        self.warnings: List[ValidationWarning] = []
        logger.info("Validate service initialized")
    
    def validate_document(
        self,
        file_path: str,
        rules: Dict[str, Any],
    ) -> ValidationResult:
        """
        校验文档
        
        Args:
            file_path: 文件路径
            rules: 校验规则
                - required_fields: 必填字段列表
                - check_links: 是否检查链接（默认False）
                - check_style: 是否检查样式（默认False）
                - expected_table_rows: 预期表格行数（可选）
                - link_timeout_sec: 链接超时时间（秒，默认3）
            
        Returns:
            ValidationResult: 校验结果
            
        Raises:
            FileNotFoundError: 文件不存在
            ValueError: 不支持的文件格式
        """
        self.errors = []
        self.warnings = []
        
        file_path_obj = Path(file_path)
        
        # 检查文件是否存在
        if not file_path_obj.exists():
            raise FileNotFoundError(f"文件不存在: {file_path}")
        
        # 读取文件内容
        try:
            if file_path_obj.suffix.lower() in ['.html', '.htm']:
                content = self._read_html(file_path)
            elif file_path_obj.suffix.lower() == '.docx':
                content = self._read_docx(file_path)
            else:
                raise ValueError(f"不支持的文件格式: {file_path_obj.suffix}")
        except Exception as e:
            logger.error(f"读取文件失败: {e}")
            raise
        
        # 1. 必填字段检查
        if rules.get("required_fields"):
            self._check_required_fields(content, rules["required_fields"])
        
        # 2. 表格行数检查
        if "expected_table_rows" in rules:
            self._check_table_rows(content, rules["expected_table_rows"])
        
        # 3. 链接有效性检查（可选）
        if rules.get("check_links", False):
            timeout = rules.get("link_timeout_sec", 3)
            self._check_links(content, timeout)
        
        # 4. 样式一致性检查（可选）
        if rules.get("check_style", False):
            self._check_style_consistency(content)
        
        return self._build_result()
    
    def check_data_alignment(
        self,
        content: str,
        data: Dict[str, Any],
    ) -> List[ValidationError]:
        """
        检查数据对齐（表格行数匹配、字段映射检查）
        
        Args:
            content: 文档内容
            data: 数据字典
            
        Returns:
            错误列表
        """
        errors = []
        
        # 检查必填字段
        if "required_fields" in data:
            for field in data["required_fields"]:
                if field not in content:
                    errors.append(ValidationError(
                        type="missing_required_field",
                        message=f"缺失必填字段: {field}",
                        field=field
                    ))
        
        # 检查表格行数
        if "expected_table_rows" in data:
            expected_rows = data["expected_table_rows"]
            if '<table' in content.lower():
                tr_count = content.lower().count('<tr')
                data_rows = max(0, tr_count - 1)  # 减去表头行
                
                if data_rows != expected_rows:
                    errors.append(ValidationError(
                        type="table_row_mismatch",
                        message=f"表格行数不匹配: 预期 {expected_rows} 行，实际 {data_rows} 行",
                        detail={
                            "expected": expected_rows,
                            "actual": data_rows
                        }
                    ))
        
        return errors
    
    def check_links(
        self,
        content: str,
        timeout: int = 3,
    ) -> List[ValidationWarning]:
        """
        检查链接有效性
        
        Args:
            content: 文档内容
            timeout: 请求超时时间（秒），默认3秒
            
        Returns:
            警告列表
        """
        warnings = []
        
        # 提取所有HTTP/HTTPS链接
        url_pattern = r'https?://[^\s<>"]+|www\.[^\s<>"]+'
        urls = re.findall(url_pattern, content)
        
        if not urls:
            return warnings
        
        # 检查每个链接（需要requests库）
        try:
            import requests
            
            for url in urls:
                try:
                    response = requests.head(url, timeout=timeout, allow_redirects=True)
                    if response.status_code >= 400:
                        warnings.append(ValidationWarning(
                            type="broken_link",
                            message=f"链接无法访问: {url}",
                            detail={"url": url, "status_code": response.status_code}
                        ))
                    elif response.elapsed.total_seconds() > timeout * 0.8:
                        warnings.append(ValidationWarning(
                            type="slow_link",
                            message=f"链接响应较慢: {url} ({response.elapsed.total_seconds():.2f}秒)",
                            detail={"url": url, "elapsed": response.elapsed.total_seconds()}
                        ))
                except Exception as e:
                    warnings.append(ValidationWarning(
                        type="broken_link",
                        message=f"链接检查失败: {url}",
                        detail={"url": url, "error": str(e)}
                    ))
        
        except ImportError:
            logger.warning("未安装requests库，跳过链接检查")
        
        return warnings
    
    def check_style_consistency(
        self,
        content: str,
    ) -> List[ValidationWarning]:
        """
        检查样式统一性（字体、页眉页脚一致性）
        
        Args:
            content: 文档内容
            
        Returns:
            警告列表
        """
        warnings = []
        
        # HTML样式检查
        if '<style' in content.lower() or 'style=' in content.lower():
            # 提取所有内联样式
            inline_styles = re.findall(r'style="([^"]+)"', content, re.IGNORECASE)
            
            # 检查字体是否一致
            font_families = set()
            for style in inline_styles:
                font_match = re.search(r'font-family:\s*([^;]+)', style, re.IGNORECASE)
                if font_match:
                    font_families.add(font_match.group(1).strip())
            
            if len(font_families) > 3:
                warnings.append(ValidationWarning(
                    type="inconsistent_fonts",
                    message=f"文档中使用了过多的字体 ({len(font_families)} 种)",
                    detail={"fonts": list(font_families)}
                ))
        
        return warnings
    
    def _read_html(self, file_path: str) -> str:
        """读取HTML文件"""
        with open(file_path, 'r', encoding='utf-8') as f:
            return f.read()
    
    def _read_docx(self, file_path: str) -> str:
        """读取Word文档内容"""
        try:
            from docx import Document
            doc = Document(file_path)
            # 提取所有段落和表格文本
            content = []
            for para in doc.paragraphs:
                content.append(para.text)
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        content.append(cell.text)
            return '\n'.join(content)
        except ImportError:
            raise ImportError("需要安装python-docx: pip install python-docx")
    
    def _check_required_fields(
        self,
        content: str,
        required_fields: List[str]
    ) -> None:
        """检查必填字段"""
        logger.info(f"检查必填字段: {required_fields}")
        
        for field in required_fields:
            if field not in content:
                self.errors.append(ValidationError(
                    type="missing_required_field",
                    message=f"缺失必填字段: {field}",
                    field=field
                ))
    
    def _check_table_rows(
        self,
        content: str,
        expected_rows: int
    ) -> None:
        """检查表格行数"""
        logger.info(f"检查表格行数，预期: {expected_rows}")
        
        # HTML表格行数检查
        if '<table' in content.lower():
            tr_count = content.lower().count('<tr')
            data_rows = max(0, tr_count - 1)  # 减去表头行
            
            if data_rows != expected_rows:
                self.warnings.append(ValidationWarning(
                    type="table_row_mismatch",
                    message=f"表格行数不匹配: 预期 {expected_rows} 行，实际 {data_rows} 行",
                    detail={
                        "expected": expected_rows,
                        "actual": data_rows
                    }
                ))
    
    def _check_links(
        self,
        content: str,
        timeout: int = 3
    ) -> None:
        """检查链接有效性"""
        logger.info(f"检查链接有效性 (超时: {timeout}秒)")
        
        warnings = self.check_links(content, timeout)
        self.warnings.extend(warnings)
    
    def _check_style_consistency(self, content: str) -> None:
        """检查样式一致性"""
        logger.info("检查样式一致性")
        
        warnings = self.check_style_consistency(content)
        self.warnings.extend(warnings)
    
    def _build_result(self) -> ValidationResult:
        """构建校验结果"""
        passed = len(self.errors) == 0
        
        return ValidationResult(
            passed=passed,
            errors=self.errors,
            warnings=self.warnings
        )

