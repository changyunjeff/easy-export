#!/usr/bin/env python3
"""
MVP: Word文档模板填充JSON/CSV数据并导出
功能：读取Word模板，使用JSON或CSV数据填充，渲染并导出为PDF（默认）或Word文档
依赖：pip install docxtpl python-docx docx2pdf
"""

import json
import csv
from pathlib import Path
from typing import Dict, Any, Optional, List
from docxtpl import DocxTemplate


class WordExporter:
    """Word模板导出器"""
    
    def _render_docx(
        self,
        template_path: str,
        data: Dict[str, Any],
        output_path: Optional[str] = None
    ) -> str:
        """
        渲染Word模板为临时docx文件
        
        Args:
            template_path: Word模板文件路径（.docx格式）
            data: 用于填充模板的JSON数据字典
            output_path: 输出docx文件路径，如果为None则创建临时文件
            
        Returns:
            临时docx文件路径
        """
        # 加载Word模板
        doc = DocxTemplate(template_path)
        
        # 渲染模板（使用Jinja2语法）
        doc.render(data)
        
        # 如果未指定输出路径，创建临时文件
        if output_path is None:
            import tempfile
            temp_file = tempfile.NamedTemporaryFile(suffix='.docx', delete=False)
            output_path = temp_file.name
            temp_file.close()
        
        # 确保输出目录存在
        output_file = Path(output_path)
        output_file.parent.mkdir(parents=True, exist_ok=True)
        
        # 保存文件
        doc.save(str(output_file))
        
        return str(output_file.absolute())
    
    def render_from_file(
        self,
        template_path: str,
        data: Dict[str, Any],
        output_path: Optional[str] = None,
        format: str = 'pdf'
    ) -> str:
        """
        从文件读取Word模板并渲染导出（默认导出为PDF）
        
        Args:
            template_path: Word模板文件路径（.docx格式）
            data: 用于填充模板的JSON数据字典
            output_path: 输出文件路径，如果为None则自动生成（基于模板文件名）
            format: 输出格式，'pdf'（默认）或 'docx'
            
        Returns:
            输出文件路径
        """
        # 如果未指定输出路径，自动生成
        if output_path is None:
            template_name = Path(template_path).stem
            output_path = f"{template_name}_output.{format}"
        
        # 确保输出路径扩展名正确
        output_file = Path(output_path)
        if output_file.suffix.lower() != f'.{format}':
            output_path = str(output_file.with_suffix(f'.{format}'))
        
        # 渲染为docx（临时文件或最终文件）
        if format.lower() == 'pdf':
            # 需要先渲染为docx，再转换为PDF
            import tempfile
            temp_docx = tempfile.NamedTemporaryFile(suffix='.docx', delete=False)
            temp_docx_path = temp_docx.name
            temp_docx.close()
            
            try:
                # 渲染为临时docx
                self._render_docx(template_path, data, temp_docx_path)
                
                # 转换为PDF
                self._convert_docx_to_pdf(temp_docx_path, output_path)
                
                return str(Path(output_path).absolute())
            finally:
                # 清理临时文件
                try:
                    Path(temp_docx_path).unlink()
                except:
                    pass
        else:
            # 直接导出为docx
            return self._render_docx(template_path, data, output_path)
    
    def _convert_docx_to_pdf(self, docx_path: str, pdf_path: str) -> None:
        """
        将docx文件转换为PDF
        
        Args:
            docx_path: docx文件路径
            pdf_path: 输出PDF文件路径
        """
        try:
            from docx2pdf import convert
            # 确保输出目录存在
            Path(pdf_path).parent.mkdir(parents=True, exist_ok=True)
            convert(docx_path, pdf_path)
        except ImportError:
            raise ImportError(
                "需要安装 docx2pdf 库才能导出PDF。\n"
                "请运行: pip install docx2pdf\n"
                "注意: docx2pdf 需要系统安装 LibreOffice 或 Microsoft Word"
            )
        except Exception as e:
            raise RuntimeError(
                f"PDF转换失败: {str(e)}\n"
                "请确保系统已安装 LibreOffice 或 Microsoft Word，"
                "并且 docx2pdf 库已正确安装。"
            )
    
def load_json_data(json_path: str) -> Dict[str, Any]:
    """
    从JSON文件加载数据
    
    Args:
        json_path: JSON文件路径
        
    Returns:
        解析后的字典数据
    """
    json_file = Path(json_path)
    if not json_file.exists():
        abs_path = json_file.absolute()
        raise FileNotFoundError(
            f"JSON文件不存在: {json_path}\n"
            f"绝对路径: {abs_path}\n"
            f"请检查文件路径是否正确。"
        )
    
    with open(json_path, 'r', encoding='utf-8') as f:
        return json.load(f)


def load_csv_data(csv_path: str) -> Dict[str, Any]:
    """
    从CSV文件加载数据并转换为字典格式
    
    支持两种CSV格式：
    1. 键值对格式（第一列是键，第二列是值）：
       key,value
       title,订单确认
       name,张三
       
    2. 表格式（第一行是表头，后续行是数据，自动转换为数组）：
       label,value
       订单号,ORD-2024-001
       金额,¥299.00
       转换为: info_items = [{'label': '订单号', 'value': 'ORD-2024-001'}, ...]
    
    Args:
        csv_path: CSV文件路径
        
    Returns:
        解析后的字典数据
    """
    csv_file = Path(csv_path)
    if not csv_file.exists():
        abs_path = csv_file.absolute()
        raise FileNotFoundError(
            f"CSV文件不存在: {csv_path}\n"
            f"绝对路径: {abs_path}\n"
            f"请检查文件路径是否正确。"
        )
    
    data: Dict[str, Any] = {}
    
    with open(csv_path, 'r', encoding='utf-8') as f:
        # 使用csv.reader解析
        reader = csv.reader(f)
        rows = list(reader)
        
        if not rows:
            return data
        
        # 检查是否是表格式（第一行是表头，且有多行数据）
        if len(rows) > 1 and len(rows[0]) >= 2:
            # 检查第一行是否是表头（通常是字符串，不是纯数字）
            first_row = rows[0]
            second_row = rows[1] if len(rows) > 1 else []
            
            # 判断：如果第一行和第二行的列数相同，且第一行看起来像表头
            if (len(first_row) == len(second_row) and 
                all(isinstance(cell, str) and not cell.replace('.', '').replace('-', '').isdigit() 
                    for cell in first_row if cell)):
                # 表格式：转换为数组
                headers = [cell.strip() for cell in first_row]
                
                # 如果表头是 label,value，则转换为 info_items 数组
                if len(headers) == 2 and 'label' in headers and 'value' in headers:
                    label_idx = headers.index('label')
                    value_idx = headers.index('value')
                    info_items = []
                    for row in rows[1:]:
                        if len(row) > max(label_idx, value_idx):
                            info_items.append({
                                'label': row[label_idx].strip() if label_idx < len(row) else '',
                                'value': row[value_idx].strip() if value_idx < len(row) else ''
                            })
                    data['info_items'] = info_items
                else:
                    # 其他表格式：转换为以第一列值为键的字典数组
                    key_header = headers[0]
                    items = []
                    for row in rows[1:]:
                        if len(row) > 0:
                            item = {}
                            for i, header in enumerate(headers):
                                if i < len(row):
                                    item[header] = row[i].strip()
                            items.append(item)
                    # 使用表头第一列作为数组名
                    array_name = f"{key_header}s" if not key_header.endswith('s') else key_header
                    data[array_name] = items
            else:
                # 键值对格式
                for row in rows:
                    if len(row) >= 2:
                        key = row[0].strip()
                        value = row[1].strip()
                        # 尝试转换数据类型
                        if value.lower() == 'true':
                            value = True
                        elif value.lower() == 'false':
                            value = False
                        elif value.isdigit():
                            value = int(value)
                        elif value.replace('.', '', 1).isdigit():
                            value = float(value)
                        data[key] = value
        else:
            # 键值对格式（只有一行或格式不明确）
            for row in rows:
                if len(row) >= 2:
                    key = row[0].strip()
                    value = row[1].strip()
                    # 尝试转换数据类型
                    if value.lower() == 'true':
                        value = True
                    elif value.lower() == 'false':
                        value = False
                    elif value.isdigit():
                        value = int(value)
                    elif value.replace('.', '', 1).isdigit():
                        value = float(value)
                    data[key] = value
    
    return data


def load_data(data_path: str) -> Dict[str, Any]:
    """
    从文件加载数据（自动检测JSON或CSV格式）
    
    Args:
        data_path: 数据文件路径（.json 或 .csv）
        
    Returns:
        解析后的字典数据
        
    Raises:
        ValueError: 不支持的文件格式
    """
    data_file = Path(data_path)
    suffix = data_file.suffix.lower()
    
    if suffix == '.json':
        return load_json_data(data_path)
    elif suffix == '.csv':
        return load_csv_data(data_path)
    else:
        raise ValueError(
            f"不支持的文件格式: {suffix}\n"
            f"支持格式: .json, .csv"
        )


def fix_table_loop_in_template(template_path: str, output_path: Optional[str] = None) -> str:
    """
    修复Word模板中的表格循环语法
    
    此函数会在表格的数据行中添加正确的Jinja2循环语法。
    使用docxtpl的底层XML操作来正确设置循环语法。
    
    Args:
        template_path: 模板文件路径
        output_path: 输出文件路径，如果为None则覆盖原文件
        
    Returns:
        修复后的文件路径
    """
    try:
        from docx import Document
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn
        import zipfile
        import xml.etree.ElementTree as ET
    except ImportError:
        print("错误: 需要安装 python-docx 库")
        print("请运行: pip install python-docx")
        return template_path
    
    if output_path is None:
        output_path = template_path
    
    # 使用docxtpl加载模板以访问XML
    doc = DocxTemplate(template_path)
    
    # 查找包含"项目"和"内容"表头的表格
    # docxtpl使用python-docx的Document对象
    for table in doc.get_docx().tables:
        if len(table.rows) > 0:
            header_cells = table.rows[0].cells
            # 检查是否是目标表格（包含"项目"和"内容"表头）
            if (len(header_cells) >= 2 and 
                header_cells[0].text.strip() == '项目' and 
                header_cells[1].text.strip() == '内容'):
                
                # 如果表格只有表头，添加一个数据行
                if len(table.rows) == 1:
                    data_row = table.add_row()
                else:
                    data_row = table.rows[1]
                
                # 在表格行的单元格中添加循环语法
                # docxtpl使用 {% tr for item in list %} 语法
                # 注意：需要在单元格的段落中正确设置
                first_cell = data_row.cells[0]
                if len(first_cell.paragraphs) == 0:
                    first_cell.add_paragraph()
                first_paragraph = first_cell.paragraphs[0]
                
                # 清空第一个单元格，添加循环语法
                first_paragraph.clear()
                # 在第一个单元格中添加循环开始和变量
                first_paragraph.add_run('{% tr for item in info_items %}')
                first_paragraph.add_run('{{ item.label }}')
                
                # 在第二个单元格中添加变量和循环结束
                second_cell = data_row.cells[1]
                if len(second_cell.paragraphs) == 0:
                    second_cell.add_paragraph()
                second_paragraph = second_cell.paragraphs[0]
                second_paragraph.clear()
                second_paragraph.add_run('{{ item.value }}')
                second_paragraph.add_run('{% endtr %}')
                
                print(f"✓ 已修复表格循环语法")
                print("  已在表格数据行中添加：{% tr for item in info_items %} 和 {% endtr %}")
                break
    
    # 保存文件
    output_file = Path(output_path)
    output_file.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_file))
    
    return str(output_file.absolute())


def create_example_template(output_path: str = 'mvp/example_template.docx'):
    """
    创建示例Word模板文件
    
    Args:
        output_path: 输出模板文件路径
    """
    try:
        from docx import Document
        from docx.shared import Pt, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
    except ImportError:
        print("错误: 需要安装 python-docx 库")
        print("请运行: pip install python-docx")
        return
    
    # 创建新文档
    doc = Document()
    
    # 设置默认字体
    style = doc.styles['Normal']
    font = style.font
    font.name = '宋体'
    font.size = Pt(12)
    
    # 标题
    title = doc.add_heading('{{ title | default("订单确认") }}', level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # 问候语
    doc.add_paragraph('亲爱的 {{ name | default("客户") }}，')
    doc.add_paragraph('')
    
    # 正文
    doc.add_paragraph('感谢您的购买！您的订单已成功提交。')
    doc.add_paragraph('')
    
    # 信息表格
    doc.add_paragraph('订单信息：')
    table = doc.add_table(rows=2, cols=2)  # 创建表头 + 一个示例数据行
    table.style = 'Light Grid Accent 1'
    
    # 表头
    header_cells = table.rows[0].cells
    header_cells[0].text = '项目'
    header_cells[1].text = '内容'
    
    # 在表格中添加一个示例行，使用Jinja2循环语法
    # 注意：docxtpl需要在表格行中使用 {% tr for %} 语法来循环添加行
    # 这里创建一个示例行，用户需要在Word中手动添加循环语法
    example_row = table.rows[1]
    example_row.cells[0].text = '{{ item.label }}'
    example_row.cells[1].text = '{{ item.value }}'
    
    # 添加说明注释（这些不会出现在最终文档中，只是提示）
    doc.add_paragraph('')
    doc.add_paragraph('注意：表格中的循环语法需要在Word中手动添加。')
    doc.add_paragraph('在表格的示例行前后添加：{% for item in info_items %} 和 {% endfor %}')
    doc.add_paragraph('或者使用：{% tr for item in info_items %} 和 {% endtr %}')
    doc.add_paragraph('')
    
    # 订单详情
    doc.add_paragraph('订单号: {{ order_id | default("ORD-2024-001234") }}')
    doc.add_paragraph('订单金额: {{ amount | default("¥299.00") }}')
    doc.add_paragraph('下单时间: {{ order_time | default("2024-01-15 14:30:00") }}')
    doc.add_paragraph('预计送达: {{ delivery_time | default("2024-01-18") }}')
    doc.add_paragraph('')
    
    # 条件判断示例
    doc.add_paragraph('{% if has_discount %}')
    doc.add_paragraph('您享受了折扣优惠！')
    doc.add_paragraph('{% endif %}')
    doc.add_paragraph('')
    
    # 页脚
    doc.add_paragraph('{{ footer_text | default("此邮件由系统自动发送，请勿直接回复。") }}')
    doc.add_paragraph('')
    doc.add_paragraph('{{ company_name | default("Easy Export") }}')
    doc.add_paragraph('{{ company_address | default("") }}')
    doc.add_paragraph('联系邮箱: {{ contact_email | default("support@example.com") }}')
    doc.add_paragraph('')
    doc.add_paragraph('© {{ current_year | default("2024") }} All rights reserved.')
    
    # 保存文件
    output_file = Path(output_path)
    output_file.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_file))
    
    print(f"✓ 示例模板已创建: {output_file.absolute()}")
    print("\n注意: 此模板文件中的Jinja2语法占位符需要手动编辑Word文档来添加。")
    print("请打开Word文档，将示例文本替换为Jinja2语法，例如：")
    print("  - 将 '订单号: ORD-2024-001234' 替换为 '订单号: {{ order_id }}'")
    print("  - 将循环部分替换为 '{% for item in info_items %}...{% endfor %}'")


def main():
    """主函数：示例用法"""
    import sys
    
    # 修复模板中的表格循环语法
    if len(sys.argv) > 1 and sys.argv[1] == '--fix-table':
        template_path = sys.argv[2] if len(sys.argv) > 2 else 'mvp/example_template.docx'
        output_path = sys.argv[3] if len(sys.argv) > 3 else None
        result_path = fix_table_loop_in_template(template_path, output_path)
        print(f"✓ 模板已修复: {result_path}")
        print("\n注意：如果自动修复不成功，请按照以下步骤手动修复：")
        print("1. 使用Microsoft Word打开模板文件")
        print("2. 找到表格的第一个数据行（表头下的第一行）")
        print("3. 在该行的第一个单元格前添加：{% tr for item in info_items %}")
        print("4. 在该行的最后一个单元格后添加：{% endtr %}")
        print("5. 确保单元格内容为：{{ item.label }} 和 {{ item.value }}")
        return
    
    # 创建示例模板
    if len(sys.argv) > 1 and sys.argv[1] == '--create-template':
        output_path = sys.argv[2] if len(sys.argv) > 2 else 'mvp/example_template.docx'
        create_example_template(output_path)
        return
    
    # 使用模板文件和数据文件（JSON或CSV）
    if len(sys.argv) >= 3:
        template_path = sys.argv[1]
        data_path = sys.argv[2]
        output_path = sys.argv[3] if len(sys.argv) > 3 else None
        
        # 从输出路径推断格式，或默认使用PDF
        if output_path:
            output_file = Path(output_path)
            if output_file.suffix.lower() == '.docx':
                format = 'docx'
            elif output_file.suffix.lower() == '.pdf':
                format = 'pdf'
            else:
                format = 'pdf'  # 默认PDF
                output_path = None  # 让函数自动生成
        else:
            format = 'pdf'  # 默认PDF
        
        # 加载数据（自动检测JSON或CSV格式）
        data = load_data(data_path)
        
        # 创建导出器
        exporter = WordExporter()
        
        # 导出（默认PDF）
        result_path = exporter.render_from_file(template_path, data, output_path, format=format)
        print(f"✓ 导出成功 ({format.upper()}): {result_path}")
        
    else:
        # 显示使用说明
        print("Word模板导出工具（默认导出PDF）")
        print("\n使用示例:")
        print("  python word_export.py <template.docx> <data.json> [output.pdf]")
        print("  python word_export.py <template.docx> <data.csv> [output.pdf]")
        print("  python word_export.py <template.docx> <data.json> output.docx  # 导出Word格式")
        print("\n创建示例模板:")
        print("  python word_export.py --create-template [output_path]")
        print("\n修复表格循环语法:")
        print("  python word_export.py --fix-table [template.docx] [output.docx]")
        print("  注意：如果表格数据未填充，使用此命令修复模板中的循环语法")
        print("\n或者使用代码:")
        print("""
from mvp.word_export import WordExporter

# 创建导出器
exporter = WordExporter()

# 准备数据（可以是字典或从JSON/CSV文件加载）
data = {
    'title': '订单确认',
    'name': '张三',
    'order_id': 'ORD-2024-001234',
    'amount': '¥299.00',
    'order_time': '2024-01-15 14:30:00',
    'delivery_time': '2024-01-18',
    'info_items': [
        {'label': '订单号', 'value': 'ORD-2024-001234'},
        {'label': '金额', 'value': '¥299.00'}
    ],
    'has_discount': True,
    'footer_text': '感谢您的使用',
    'company_name': 'Easy Export',
    'contact_email': 'support@example.com',
    'current_year': 2024
}

# 或从文件加载
from mvp.word_export import load_data
data = load_data('data.json')  # 或 load_data('data.csv')

# 导出为PDF（默认）
exporter.render_from_file('template.docx', data, 'output.pdf')

# 或导出为Word
exporter.render_from_file('template.docx', data, 'output.docx', format='docx')
        """)


if __name__ == '__main__':
    main()

