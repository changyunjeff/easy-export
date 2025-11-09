#!/usr/bin/env python3
"""
创建Word模板文件的辅助脚本
功能：生成一个包含Jinja2语法占位符的Word模板文件
依赖：pip install python-docx
"""

from pathlib import Path


def create_word_template(output_path: str = 'mvp/example_template.docx'):
    """
    创建示例Word模板文件（包含Jinja2语法占位符）
    
    注意：由于python-docx无法直接创建包含Jinja2语法的文档，
    此脚本会创建一个基础结构，然后需要手动在Word中编辑添加Jinja2语法。
    
    Args:
        output_path: 输出模板文件路径
    """
    try:
        from docx import Document
        from docx.shared import Pt, RGBColor, Inches
        from docx.enum.text import WD_ALIGN_PARAGRAPH
    except ImportError:
        print("错误: 需要安装 python-docx 库")
        print("请运行: pip install python-docx")
        return False
    
    # 创建新文档
    doc = Document()
    
    # 设置默认字体
    style = doc.styles['Normal']
    font = style.font
    font.name = '宋体'
    font.size = Pt(12)
    
    # 标题（提示用户需要替换为Jinja2语法）
    title = doc.add_heading('【请替换为】{{ title | default("订单确认") }}', level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # 添加说明
    doc.add_paragraph('')
    note_para = doc.add_paragraph('注意：此文档中的【请替换为】标记需要手动删除，')
    note_para.add_run('并将示例文本替换为Jinja2语法占位符。').bold = True
    doc.add_paragraph('')
    doc.add_paragraph('=' * 50)
    doc.add_paragraph('')
    
    # 问候语
    doc.add_paragraph('【请替换为】亲爱的 {{ name | default("客户") }}，')
    doc.add_paragraph('')
    
    # 正文
    doc.add_paragraph('感谢您的购买！您的订单已成功提交。')
    doc.add_paragraph('')
    
    # 订单信息标题
    doc.add_paragraph('订单信息：')
    doc.add_paragraph('')
    
    # 订单详情（使用表格格式）
    table = doc.add_table(rows=5, cols=2)
    table.style = 'Light Grid Accent 1'
    
    # 填充表格数据（这些需要替换为Jinja2语法）
    rows_data = [
        ('订单号', '【请替换为】{{ order_id }}'),
        ('订单金额', '【请替换为】{{ amount }}'),
        ('下单时间', '【请替换为】{{ order_time }}'),
        ('预计送达', '【请替换为】{{ delivery_time }}'),
        ('配送地址', '【请替换为】{{ delivery_address | default("") }}'),
    ]
    
    for i, (label, value) in enumerate(rows_data):
        row = table.rows[i]
        row.cells[0].text = label
        row.cells[1].text = value
    
    doc.add_paragraph('')
    
    # 详细信息列表
    doc.add_paragraph('详细信息：')
    doc.add_paragraph('【请替换为以下Jinja2循环语法】')
    doc.add_paragraph('{% for item in info_items %}')
    doc.add_paragraph('{{ item.label }}: {{ item.value }}')
    doc.add_paragraph('{% endfor %}')
    doc.add_paragraph('')
    
    # 条件判断示例
    doc.add_paragraph('【请替换为以下Jinja2条件语法】')
    doc.add_paragraph('{% if has_discount %}')
    doc.add_paragraph('您享受了折扣优惠！')
    doc.add_paragraph('{% endif %}')
    doc.add_paragraph('')
    
    # 分隔线
    doc.add_paragraph('-' * 50)
    doc.add_paragraph('')
    
    # 页脚信息
    doc.add_paragraph('【请替换为】{{ footer_text | default("此邮件由系统自动发送，请勿直接回复。") }}')
    doc.add_paragraph('')
    doc.add_paragraph('【请替换为】{{ company_name | default("Easy Export") }}')
    doc.add_paragraph('【请替换为】{{ company_address }}')
    doc.add_paragraph('联系邮箱: 【请替换为】{{ contact_email | default("support@example.com") }}')
    doc.add_paragraph('')
    doc.add_paragraph('【请替换为】© {{ current_year | default("2024") }} All rights reserved.')
    
    # 保存文件
    output_file = Path(output_path)
    output_file.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_file))
    
    print(f"✓ 基础模板已创建: {output_file.absolute()}")
    print("\n下一步操作：")
    print("1. 使用Microsoft Word打开该文件")
    print("2. 删除所有【请替换为】标记")
    print("3. 将示例文本替换为对应的Jinja2语法占位符")
    print("4. 保存文件")
    print("\n参考文档: mvp/word_template_guide.md")
    
    return True


def main():
    """主函数"""
    import sys
    
    output_path = sys.argv[1] if len(sys.argv) > 1 else 'mvp/example_template.docx'
    
    if create_word_template(output_path):
        print("\n模板创建完成！")
    else:
        print("\n模板创建失败，请检查依赖是否已安装。")
        print("安装命令: pip install python-docx")


if __name__ == '__main__':
    main()

